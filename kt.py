# # Import Module
import os
import pandas as pd
from docx import Document
source = {}

# Folder Path
path = "D:\\sakthi\\kalisamy Tex\\mar-Jun23\\tds"
allDate = []
allSalary = []
expectedTotal = {}
calculatedTotal = {}

# # Change the directory
os.chdir(path)

# # Read text File
overAllTDS = {}


def read_text_file(file_path, a):

    wordDoc = Document(
        file_path)
    # data = [[cell.text for cell in row.cells] for row in wordDoc.tables.rows]
    # df = pd.DataFrame(data)
    salt = a.split(' ')[0]
    if salt not in allSalary:
        allSalary.append(salt)
    from datetime import datetime

    date_object = ""
    date_format = "%d%m%Y"
    date_string = a.split(' ')[1].split('.')[0]
    try:
        date_object = datetime.strptime(date_string, date_format)
        if date_object not in allDate:
            allDate.append(date_object)
        # print("Converted Date:", date_object.strftime("%Y-%m-%d"))
    except ValueError:
        print("Invalid date format. Please use ddmmyyyy format.")
        return -1
    date = date_object.strftime("%d-%b-%Y")

    for table in wordDoc.tables:
        for row in table.rows:
            if row.cells[3].text == "AMOUNT":
                continue
            if row.cells[2].text == "TOTAL":
                if salt not in expectedTotal:
                    expectedTotal[salt] = {}
                expectedTotal[salt][date] = int(row.cells[3].text)
                continue
            # if not a[row.cells[2].text]:
            if salt not in source:
                source[salt] = {}
            if len(row.cells[2].text) < 5:
                print(
                    "the account number is less than 5 character verify "+salt+' '+date)
            if row.cells[2].text not in source[salt]:
                source[salt][row.cells[2].text] = {}
                source[salt][row.cells[2].text]["NAME"] = row.cells[1].text
                source[salt][row.cells[2].text]["ACCOUNT NUMBER"] = row.cells[2].text
                source[salt][row.cells[2].text]["Total TDS"] = 0

            source[salt][row.cells[2].text][date] = int(row.cells[3].text)
            source[salt][row.cells[2].text]["TDS " +
                                            date] = int(row.cells[3].text)/99

            source[salt][row.cells[2].text]["Total TDS"] += int(
                row.cells[3].text)/99
            if salt not in calculatedTotal:
                calculatedTotal[salt] = {}
            if date not in calculatedTotal[salt]:
                calculatedTotal[salt][date] = 0
            calculatedTotal[salt][date] += int(row.cells[3].text)

       


# iterate through all file
for file in os.listdir():
    # Check whether file is in text format or not
    if file.endswith(".docx"):
        print(file)
        file_path = f"{path}\{file}"

        # call read text file function
        read_text_file(file_path, file)
for i in calculatedTotal:
    for j in calculatedTotal[i]:
        if "TOTAL" not in source[i]:
            source[i]["TOTAL"] = {}
            source[i]["TOTAL"]["ACCOUNT NUMBER"] = "TOTAL"
        source[i]["TOTAL"][j] = calculatedTotal[i][j]
        source[i]["TOTAL"]["TDS " + j] = calculatedTotal[i][j]/99
        if i not in overAllTDS:
            overAllTDS[i] = 0
        overAllTDS[i] += calculatedTotal[i][j]/99
for k in overAllTDS:
    source[k]["TOTAL"]["Total TDS"] = overAllTDS[k]
for m in expectedTotal:
    for n in expectedTotal[m]:
        if expectedTotal[m][n] != calculatedTotal[m][n]:
            print("Total differs at " + m + ' '+n)
allDate.sort()
date_string_array = [date_obj.strftime("%d-%b-%Y") for date_obj in allDate]

clm = ["NAME", "ACCOUNT NUMBER"]
for i in date_string_array:
    clm.append(i)
for i in date_string_array:
    clm.append("TDS "+i)
clm.append("Total TDS")
# print(source)
df1 = pd.DataFrame(source["SALARY1"])
df1.fillna(0, inplace=True)
df1T = df1.T
df1T.index = [i for i in range(1, df1T.shape[0]+1)]

df2 = pd.DataFrame(source["SALARY2"])
df2.fillna(0, inplace=True)
df2T = df2.T
df2T.index = [i for i in range(1, df2T.shape[0]+1)]

with pd.ExcelWriter(f"{path}\output.xlsx") as writer:  # doctest: +SKIP
    df1T.to_excel(writer, sheet_name='Cheque 1 Party', columns=clm)
    df2T.to_excel(writer, sheet_name='Cheque 2 Party', columns=clm)
